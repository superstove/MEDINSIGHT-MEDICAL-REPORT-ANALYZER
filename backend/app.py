# Import necessary libraries
from flask import Flask, request, jsonify, session # Import session
from flask_cors import CORS
from dotenv import load_dotenv
import os
import logging
import google.generativeai as genai
import pandas as pd
import docx
import pytesseract
import cv2
import fitz  # PyMuPDF
import pdfplumber
import pathlib
import json
import re
import numpy as np
from io import BytesIO
from PIL import Image
import secrets # Import secrets for generating a key
from googletrans import Translator # For translation
import time # Potentially for translation delays
from flask_pymongo import PyMongo
from bson import ObjectId
import jwt
import datetime
from functools import wraps
from werkzeug.security import generate_password_hash, check_password_hash
import firebase_admin
from firebase_admin import credentials, firestore
import docx2txt
import io

# --- Constants ---
MAX_CHAT_HISTORY_TURNS = 10 # Number of conversation pairs (user + model) to keep
MAX_CONTEXT_CHARS_FOR_CHAT = 30000 # Max characters of document context to inject into chat prompt
MAX_TEXT_ANALYSIS_CHARS = 200000 # Max characters fed into initial text analysis

# --- Setup ---

# Load environment variables
load_dotenv()

# Initialize Flask
app = Flask(__name__)
CORS(app, supports_credentials=True) # Allow credentials for session cookies

# --- IMPORTANT: Add a Secret Key for Session Management ---
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY', secrets.token_hex(16))
if app.config['SECRET_KEY'] == secrets.token_hex(16): # Check if using the default generated one
    logging.warning("Using a default generated FLASK_SECRET_KEY. Set a persistent FLASK_SECRET_KEY environment variable for production.")

# Logging setup
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
log = logging.getLogger('werkzeug') # Flask's default logger
log.setLevel(logging.INFO)

# Set the path to Tesseract OCR (Ensure this is correct for your system)
TESSERACT_PATH = os.getenv("TESSERACT_CMD", r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe") # Default to Windows path if not set in .env
try:
    if os.path.exists(TESSERACT_PATH):
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
        logging.info(f"Using Tesseract at: {pytesseract.pytesseract.tesseract_cmd}")
    else:
        logging.warning(f"Tesseract path '{TESSERACT_PATH}' does not exist. OCR will likely fail. Set TESSERACT_CMD environment variable if needed.")
except Exception as e:
     logging.error(f"Error setting up Tesseract path: {e}. Please ensure Tesseract is installed and the path/environment variable is correct.")

# Configure API Keys
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("❌ GEMINI_API_KEY is missing in .env file.")

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)
try:
    # Use a model capable of multimodal input and good chat performance
    gemini_model = genai.GenerativeModel("gemini-1.5-flash-latest")
    logging.info("Using gemini-1.5-flash-latest for analysis and chat.")
except Exception as e:
    logging.warning(f"Could not initialize gemini-1.5-flash-latest, falling back to gemini-pro. Error: {e}")
    gemini_model = genai.GenerativeModel("gemini-pro") # Fallback

# Uploads directory
UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Initialize Translator
translator = Translator()

# Initialize Firebase Admin
cred = credentials.Certificate('serviceAccountKey.json')  # Updated path to the JSON file in the backend directory
firebase_admin.initialize_app(cred)
db = firestore.client()

# JWT configuration
JWT_SECRET_KEY = os.getenv("JWT_SECRET_KEY")
JWT_ACCESS_TOKEN_EXPIRES = datetime.timedelta(days=1)

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        if 'Authorization' in request.headers:
            token = request.headers['Authorization'].split(' ')[1]
        
        if not token:
            return jsonify({'message': 'Token is missing'}), 401
        
        try:
            data = jwt.decode(token, JWT_SECRET_KEY, algorithms=["HS256"])
            current_user = db.collection('users').document(data['user_id']).get()
            if not current_user.exists:
                return jsonify({'message': 'User not found'}), 401
        except:
            return jsonify({'message': 'Token is invalid'}), 401
        
        return f(current_user, *args, **kwargs)
    return decorated

# User profile endpoints
@app.route('/api/user/profile', methods=['GET'])
@token_required
def get_user_profile(current_user):
    user_data = current_user.to_dict()
    # Remove sensitive information
    user_data.pop('password', None)
    user_data['id'] = current_user.id
    
    return jsonify(user_data), 200

@app.route('/api/user/profile', methods=['PUT'])
@token_required
def update_user_profile(current_user):
    try:
        data = request.json
        
        # Update only provided fields
        update_data = {}
        
        if 'name' in data:
            update_data['name'] = data['name']
        if 'personalInfo' in data:
            update_data['personalInfo'] = data['personalInfo']
        if 'contactInfo' in data:
            update_data['contactInfo'] = data['contactInfo']
            
        # Update user document
        current_user.reference.update(update_data)
        
        return jsonify({'message': 'Profile updated successfully'}), 200
    except Exception as e:
        logging.error(f"Error updating profile: {str(e)}")
        return jsonify({'message': 'Error updating profile'}), 500

@app.route('/api/user/avatar', methods=['POST'])
@token_required
def update_avatar(current_user):
    try:
        data = request.json
        
        if not data or 'avatar' not in data:
            return jsonify({'message': 'No avatar provided'}), 400
            
        # Update avatar URL
        current_user.reference.update({
            'avatar': data['avatar']
        })
        
        return jsonify({'message': 'Avatar updated successfully'}), 200
    except Exception as e:
        logging.error(f"Error updating avatar: {str(e)}")
        return jsonify({'message': 'Error updating avatar'}), 500

# -------------------- Utility Functions (Mostly Unchanged) --------------------
# (preprocess_image, extract_text_from_image, extract_text,
# parse_gemini_json_response, handle_gemini_error_response,
# analyze_text_with_gemini, analyze_image_with_gemini, translate_analysis)
# Keep the functions from the previous version here.
# ... (Include all the utility functions from the previous answer here) ...

def preprocess_image(image):
    """Preprocesses an image for better OCR."""
    try:
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        # Apply median blur to remove noise
        gray = cv2.medianBlur(gray, 3)
        # Apply thresholding - OTSU is good for bimodal histograms
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        # Optional: Dilation and Erosion to remove noise (use carefully)
        # kernel = np.ones((1, 1), np.uint8)
        # thresh = cv2.dilate(thresh, kernel, iterations=1)
        # thresh = cv2.erode(thresh, kernel, iterations=1)
        return thresh
    except cv2.error as cv_err:
         if "color conversion" in str(cv_err):
             logging.warning(f"Image seems to be already grayscale or has unsupported channels. Trying direct Otsu threshold.")
             try:
                 # If it's already gray or has alpha, try just thresholding
                 if len(image.shape) == 2: # Grayscale
                    gray = image
                 elif image.shape[2] == 4: # BGRA/RGBA
                    gray = cv2.cvtColor(image, cv2.COLOR_BGRA2GRAY)
                 else: # Unexpected shape
                    logging.error(f"Unsupported image channel count: {image.shape}")
                    return None
                 thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                 return thresh
             except Exception as fallback_err:
                 logging.error(f"Error during fallback preprocessing: {fallback_err}")
                 return None
         else:
             logging.error(f"Error during image preprocessing (OpenCV): {cv_err}")
             return None
    except Exception as e:
        logging.error(f"General error during image preprocessing: {e}")
        return None

def extract_text_from_image(file_path):
    """Extracts text from an image using OCR with preprocessing."""
    try:
        image = cv2.imread(file_path)
        if image is None:
            logging.error(f"Error: Could not read image at {file_path} using OpenCV.")
            # Try with PIL as a fallback for formats OpenCV might struggle with
            try:
                pil_image = Image.open(file_path)
                image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR) # Convert PIL to OpenCV
                logging.info(f"Successfully read image {os.path.basename(file_path)} using PIL fallback.")
            except Exception as pil_err:
                logging.error(f"PIL fallback also failed for {file_path}: {pil_err}")
                return ""

        preprocessed_image = preprocess_image(image)
        if preprocessed_image is None:
             logging.warning(f"Preprocessing failed for image {file_path}, attempting OCR on original image.")
             # Attempt OCR on the original grayscale version if preprocessing failed
             try:
                 if len(image.shape) == 3:
                     gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                 else: # Assume already gray
                     gray_image = image
                 text = pytesseract.image_to_string(gray_image, config='--oem 3 --psm 3').strip() # Use default PSM 3
                 if text:
                     logging.info(f"OCR successful on original grayscale image for {file_path} after preprocessing failure.")
                 else:
                     logging.warning(f"OCR yielded no text on original grayscale image either for {file_path}.")
                 return text
             except Exception as ocr_fallback_err:
                 logging.error(f"Error during OCR fallback attempt for {file_path}: {ocr_fallback_err}")
                 return ""

        # Try PSM 6 first (Assume a single uniform block of text)
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(preprocessed_image, config=custom_config).strip()

        if not text or len(text) < 20: # If PSM 6 yields little/no text, try PSM 3 (Auto page segmentation)
             logging.warning(f"OCR with PSM 6 yielded minimal/no text for {os.path.basename(file_path)} (len: {len(text)}), trying PSM 3 (Auto).")
             custom_config = r'--oem 3 --psm 3'
             text = pytesseract.image_to_string(preprocessed_image, config=custom_config).strip()

        if not text:
            logging.warning(f"Final OCR attempt yielded no text for {os.path.basename(file_path)}.")

        return text
    except pytesseract.TesseractNotFoundError:
         logging.error("TESSERACT NOT FOUND. Please install Tesseract and ensure pytesseract.pytesseract.tesseract_cmd points to the executable.")
         return "Error: Tesseract not found."
    except Exception as e:
        logging.error(f"Error during OCR for {os.path.basename(file_path)}: {e}", exc_info=True)
        return ""

def extract_text(file_path):
    """Extract text from PDFs, DOCX, Images (using OCR), and Excel files."""
    extracted_text = ""
    ext = pathlib.Path(file_path).suffix.lower()
    filename = os.path.basename(file_path)

    try:
        logging.info(f"Attempting to extract text from: {filename} (type: {ext})")
        if ext == ".pdf":
            text_found = False
            # Try PyMuPDF first (usually faster and better)
            try:
                doc = fitz.open(file_path)
                # Use flags for better layout preservation if needed: page.get_text("text", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)
                extracted_text = "\n".join([page.get_text("text", sort=True).strip() for page in doc if page.get_text("text").strip()]) # Added sort=True
                doc.close()
                if extracted_text.strip():
                    logging.info(f"PyMuPDF extraction successful for {filename}.")
                    text_found = True
                else:
                    logging.warning(f"PyMuPDF extracted no text from {filename}. It might be image-based or empty.")
            except Exception as e_fitz:
                logging.warning(f"PyMuPDF failed for {filename}: {e_fitz}. Trying pdfplumber...")

            # Try pdfplumber if PyMuPDF failed or got no text
            if not text_found:
                try:
                    with pdfplumber.open(file_path) as pdf:
                        # Adjust tolerances if text extraction is poor on certain PDFs
                        extracted_text = "\n".join([page.extract_text(x_tolerance=1, y_tolerance=3) or "" for page in pdf.pages]) # Increased y_tolerance slightly
                    if extracted_text.strip():
                        logging.info(f"pdfplumber extraction successful for {filename}.")
                        text_found = True
                    else:
                        logging.warning(f"pdfplumber also extracted no text from {filename}.")
                except Exception as e_plumber:
                    logging.error(f"Both PyMuPDF and pdfplumber failed for PDF {filename}. Last error: {e_plumber}")
                    extracted_text = "" # Ensure it's empty if both fail

            # Future Improvement: Add OCR step here if text_found is still False for image-based PDFs

        elif ext == ".docx":
            try:
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                # Optional: Add table extraction
                # for table in doc.tables:
                #     for row in table.rows:
                #         row_text = [cell.text for cell in row.cells]
                #         full_text.append("\t".join(row_text)) # Tab-separated cells
                extracted_text = "\n".join(filter(None, full_text)) # Join non-empty parts
                logging.info(f"Docx extraction successful for {filename}.")
            except Exception as e_docx:
                 logging.error(f"Error extracting from DOCX {filename}: {e_docx}")
                 extracted_text = ""


        elif ext == ".xlsx":
            try:
                # Read all sheets, convert to string, handling potential errors
                excel_file = pd.ExcelFile(file_path)
                all_text = []
                for sheet_name in excel_file.sheet_names:
                    try:
                        df = excel_file.parse(sheet_name)
                        # Convert dataframe to string, replacing NaN with empty string
                        df_string = df.fillna('').to_string(index=False, header=True)
                        all_text.append(f"--- Sheet: {sheet_name} ---\n{df_string}")
                    except Exception as e_sheet:
                        logging.warning(f"Could not read sheet '{sheet_name}' in {filename}: {e_sheet}")
                        all_text.append(f"--- Sheet: {sheet_name} (Error reading) ---")
                extracted_text = "\n\n".join(all_text)
                logging.info(f"Xlsx extraction successful for {filename}.")
            except Exception as e_xlsx:
                logging.error(f"Error reading XLSX {filename}: {e_xlsx}")
                extracted_text = ""

        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]:
            extracted_text = extract_text_from_image(file_path)
            # Logging is done inside extract_text_from_image

        else:
            logging.warning(f"Unsupported file type for standard text extraction: {ext}")
            # Fallback: Try reading as plain text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
                if extracted_text.strip():
                    logging.info(f"Read {filename} as plain text (fallback).")
                else:
                    extracted_text = ""
            except Exception as e_txt:
                logging.warning(f"Could not read {filename} as plain text: {e_txt}")
                extracted_text = ""

        # Log length and snippet
        text_length = len(extracted_text)
        logging.info(f"📄 Extracted Text Length for {filename}: {text_length} chars")
        if text_length > 0:
             logging.debug(f"📄 Extracted Text Snippet (first 300 chars): {extracted_text[:300]}...")
        elif ext not in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]: # Avoid logging no text for images if OCR failed
            logging.warning(f"No text could be extracted from {filename}.")


    except Exception as e:
        logging.error(f"❌ Unexpected error extracting text from {filename}: {e}", exc_info=True)
        return "" # Return empty string on error

    return extracted_text.strip()

def parse_gemini_json_response(response_text, context="analysis"):
    """Attempts to parse JSON from Gemini response, handling markdown code blocks."""
    logging.debug(f"Raw Gemini response for {context}:\n---\n{response_text}\n---")
    # More robust regex to handle potential variations in markdown/spacing
    json_match = re.search(r"```(?:json)?\s*(\{.*\})\s*```|^\s*(\{.*\})\s*$", response_text, re.DOTALL | re.IGNORECASE | re.MULTILINE)
    if json_match:
        json_string = json_match.group(1) or json_match.group(2) # Group 1 for markdown, Group 2 for direct object
        if json_string:
            try:
                structured_output = json.loads(json_string)
                logging.info(f"✅ Structured AI {context} parsed successfully (from regex match).")
                return structured_output
            except json.JSONDecodeError as e:
                logging.error(f"❌ AI {context} response failed JSON parsing (after regex): {e}\n--- JSON String Attempted ---\n{json_string}\n---")
                return {"error": f"AI {context} response could not be parsed as JSON (invalid structure)", "raw_text_snippet": response_text[:500]}
        else:
             # This case should be rare with the improved regex but handled just in case
             logging.error(f"❌ Regex matched but failed to extract JSON content ({context}). Raw: {response_text[:500]}...")
             return {"error": f"Regex matched but failed to extract JSON content for {context}.", "raw_text_snippet": response_text[:500]}
    else:
        # Try parsing directly if no markdown block found or the regex didn't catch it
        try:
            # Attempt to find the first '{' and last '}' to handle potential leading/trailing text
            start_index = response_text.find('{')
            end_index = response_text.rfind('}')
            if start_index != -1 and end_index != -1 and end_index > start_index:
                potential_json = response_text[start_index : end_index + 1]
                try:
                    structured_output = json.loads(potential_json)
                    logging.info(f"✅ Structured AI {context} parsed successfully (direct parse with slicing).")
                    return structured_output
                except json.JSONDecodeError:
                    logging.warning(f"Direct parse with slicing failed for {context}. Trying full direct parse.")
                    # Fall through to try parsing the whole string if slicing fails

            # Try parsing the whole response text directly
            structured_output = json.loads(response_text)
            logging.info(f"✅ Structured AI {context} parsed successfully (direct full parse).")
            return structured_output
        except json.JSONDecodeError as e_direct:
             logging.error(f"❌ No JSON block found and direct/sliced parse failed ({context}): {e_direct}")
             return {"error": f"No valid JSON found in AI {context} response", "raw_text_snippet": response_text[:500]}

def handle_gemini_error_response(response, context="analysis"):
    """Logs details and creates error dict for failed Gemini responses."""
    logging.error(f"❌ Gemini {context} generation failed or was blocked.")
    error_payload = {"error": f"AI {context} failed with no detailed reason provided."} # Default
    try:
        # Attempt to get feedback/reason if available (new API structure)
        if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
            block_reason = getattr(response.prompt_feedback, 'block_reason', 'Unknown')
            safety_ratings = getattr(response.prompt_feedback, 'safety_ratings', 'N/A')
            safety_ratings_str = json.dumps(safety_ratings) if isinstance(safety_ratings, list) else str(safety_ratings) # Format ratings nicely
            logging.error(f"Prompt Feedback: Block Reason: {block_reason}, Safety Ratings: {safety_ratings_str}")
            error_payload = {"error": f"AI {context} request was blocked. Reason: {block_reason}", "safety_ratings": safety_ratings}
        # Check candidates finish reason if feedback isn't the primary source
        elif hasattr(response, 'candidates') and response.candidates:
             # Check the first candidate
             candidate = response.candidates[0]
             finish_reason = getattr(candidate, 'finish_reason', 'UNKNOWN')
             if finish_reason != 'STOP': # If it didn't finish normally
                 safety_ratings = getattr(candidate, 'safety_ratings', 'N/A')
                 safety_ratings_str = json.dumps(safety_ratings) if isinstance(safety_ratings, list) else str(safety_ratings)
                 logging.error(f"Candidate Finish Reason: {finish_reason}, Safety Ratings: {safety_ratings_str}")
                 error_payload = {"error": f"AI {context} generation stopped unexpectedly. Reason: {finish_reason}", "safety_ratings": safety_ratings}
             else:
                 # It stopped, but maybe content was missing? (Handled elsewhere, but log here)
                 logging.warning(f"Gemini {context} stopped normally but content might be missing or failed parsing.")
                 error_payload = {"error": f"AI {context} response processing failed after generation."}
        else:
             logging.error("No candidates or prompt feedback found in the Gemini error response.")

    except Exception as e_feedback:
        logging.error(f"Error trying to access error details from Gemini response: {e_feedback}")
        error_payload = {"error": f"AI {context} failed, and error details could not be extracted."}

    return error_payload


def analyze_text_with_gemini(text):
    """Generate a structured medical summary from text using Gemini AI."""
    try:
        if len(text) > MAX_TEXT_ANALYSIS_CHARS:
            logging.warning(f"Input text length ({len(text)} chars) exceeds limit ({MAX_TEXT_ANALYSIS_CHARS}), truncating.")
            text = text[:MAX_TEXT_ANALYSIS_CHARS] + "\n... [Content Truncated]"

        if not text.strip():
            logging.warning("analyze_text_with_gemini called with empty text.")
            return {"error": "Input text was empty."}

        logging.info("🔄 Sending text to Gemini for structured analysis...")

        prompt = f"""
        Analyze the following medical report text and extract the information into a structured JSON format.
        Be comprehensive but concise. If information for a field is not present, use `null` or an empty list `[]`.
        Focus *only* on the information present in the text provided. Do not infer or add external knowledge.

        YOUR RESPONSE MUST BE A SINGLE, VALID JSON OBJECT AND NOTHING ELSE.
        DO NOT include any explanatory text before or after the JSON object.
        DO NOT use markdown formatting like ```json.

        JSON Format:
        {{
            "summary": "A concise summary of the main points of the medical report excerpt. Synthesize key findings.",
            "diagnosis": "Primary or potential diagnosis mentioned, or null.",
            "key_findings": ["List key observations, test results, or findings mentioned.", "Finding 2", "..."],
            "causes": ["Possible causes mentioned for the condition, or null.", "Cause 2", "..."],
            "recommendations": "Specific medical recommendations or next steps mentioned, or null.",
            "precautions": ["Any precautions advised in the text.", "Precaution 2", "..."],
            "remedies": ["Mentioned treatments, therapies, or remedies.", "Remedy 2", "..."],
            "important_notes": "Other significant details or notes from the report.",
            "treatment_plan": "Outline of the treatment plan if described, or null.",
            "lifestyle_changes": ["Specific lifestyle changes suggested.", "Change 2", "..."],
            "urgent_concerns": "Any explicitly mentioned urgent concerns or red flags, or null."
        }}

        Medical Report Text:
        --- START TEXT ---
        {text}
        --- END TEXT ---

        Now, provide ONLY the JSON object based on the text above.
        """

        try:
             safety_settings = { # Adjust safety settings as needed for medical content
                'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE',
                'HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE',
                'HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_LOW_AND_ABOVE', # Might need BLOCK_NONE if issues arise
                'HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE',
             }
             response = gemini_model.generate_content(prompt, safety_settings=safety_settings)

             # Check for valid response content
             if response.candidates and response.candidates[0].content and response.candidates[0].content.parts:
                 response_text = response.candidates[0].content.parts[0].text.strip()
                 return parse_gemini_json_response(response_text, context="text analysis")
             else:
                 # Handle cases where generation failed or was blocked
                 return handle_gemini_error_response(response, context="text analysis")

        except Exception as api_err:
             logging.exception("❌ Error during Gemini API call for text analysis:")
             if "API key not valid" in str(api_err):
                 return {"error": "Gemini API key is invalid. Please check your .env file."}
             return {"error": f"Gemini API communication error: {api_err}"}


    except Exception as e:
        logging.exception("❌ Unexpected error during Gemini text analysis setup:")
        return {"error": f"Unexpected AI text processing error: {e}"}


def analyze_image_with_gemini(image_data, filename="image"):
    """Analyzes an image using Gemini AI and returns a structured JSON."""
    try:
        logging.info(f"🔄 Preparing image '{filename}' for Gemini analysis...")
        prompt = """
        YOUR **ONLY** OUTPUT MUST BE A VALID JSON OBJECT.
        DO NOT INCLUDE ANY INTRODUCTORY TEXT, EXPLANATIONS, NOTES, OR MARKDOWN FORMATTING (like ```json) BEFORE OR AFTER THE JSON OBJECT.

        The JSON object MUST conform EXACTLY to the following format. If you cannot determine a value for a particular field based *only* on the visual information in the image, set it to null or an empty list [] if it's an array. Do not infer information not visually present.

        JSON Format:
        {
            "summary": "Provide a DETAILED and COMPREHENSIVE summary explaining the primary medical issue or abnormality visible in the image. Synthesize the key findings into a coherent description of the problem. Include observations about location, extent, appearance, and relationship to surrounding structures.",
            "diagnosis": "Potential diagnosis based *only* on the visual evidence in the image, if possible. State if a diagnosis cannot be determined from the image alone. Set to null if no specific visual signs suggest a diagnosis.",
            "key_findings": ["Finding 1", "Finding 2", "List all significant visual observations, describing what you see."],
            "precautions": ["General precautions potentially relevant ONLY to the visual findings (e.g., 'potential fracture site immobilization'), or null.", "Precaution 2"],
            "remedies": ["Potential remedies/treatments suggested ONLY by the visual findings (e.g., 'possible need for drainage if fluid collection confirmed'), or null.", "Remedy 2"],
            "urgent_concerns": "Any visually evident urgent concern (e.g., 'significant mass effect', 'pneumothorax', 'active hemorrhage'), or null.",
            "anatomical_structures": ["Structure 1", "List main visible anatomical structures identified in the image."]
        }

        Analyze the provided medical image and generate ONLY the JSON output based on the fields above. Ensure the 'summary' and 'key_findings' fields are thorough and based solely on the visual content.
        """

        try:
            image = Image.open(BytesIO(image_data))
            logging.info(f"Image '{filename}' loaded via PIL. Original format: {image.format}, Mode: {image.mode}, Size: {image.size}")

            # Ensure image is RGB
            if image.mode == 'RGBA':
                 bg = Image.new('RGB', image.size, (255, 255, 255))
                 bg.paste(image, (0, 0), image)
                 image = bg
                 logging.info("Converted RGBA image to RGB by pasting on white background.")
            elif image.mode != 'RGB':
                 logging.info(f"Converting image mode from {image.mode} to RGB.")
                 image = image.convert('RGB')

            buffer = BytesIO()
            image.save(buffer, format='JPEG', quality=90)
            prepared_image_data = buffer.getvalue()
            image_mime = 'image/jpeg'
            logging.info(f"Image '{filename}' prepared for API as {image_mime}. Size: {len(prepared_image_data)} bytes.")

        except Exception as img_err:
            logging.error(f"❌ Failed to load/process image '{filename}' with PIL: {img_err}", exc_info=True)
            return {"error": f"Failed to process image data for {filename}: {img_err}"}

        content = [ # Changed to list format for clarity, though single image+prompt works too
            {"mime_type": image_mime, "data": prepared_image_data},
            {"text": prompt}
        ]

        try:
            logging.info(f"🔄 Sending image '{filename}' to Gemini for visual analysis...")
            safety_settings = { # Adjust safety settings for medical images
                'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE',
                'HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE',
                'HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_LOW_AND_ABOVE', # Be cautious, may need BLOCK_NONE
                'HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE',
            }
            # Use generate_content directly with the list of parts
            response = gemini_model.generate_content(content, safety_settings=safety_settings)
            logging.info(f"Received response object from Gemini for image '{filename}'.")

            # Check for valid response content
            if response.candidates and response.candidates[0].content and response.candidates[0].content.parts:
                response_text = response.candidates[0].content.parts[0].text.strip()
                return parse_gemini_json_response(response_text, context=f"image analysis ({filename})")
            else:
                # Handle cases where generation failed or was blocked
                return handle_gemini_error_response(response, context=f"image analysis ({filename})")

        except Exception as api_err:
            logging.exception(f"❌ Error during Gemini API call for image analysis ({filename}):")
            if "API key not valid" in str(api_err):
                 return {"error": "Gemini API key is invalid. Please check your .env file."}
            # Add specific API error checks if needed
            return {"error": f"Gemini API communication error during image analysis: {api_err}"}


    except Exception as e:
        logging.exception(f"❌ Unexpected error during Gemini image analysis setup ({filename}):")
        return {"error": f"Unexpected AI image processing error: {e}"}


def translate_analysis(analysis_dict, target_language):
    """Translates the values and optionally keys of an analysis dictionary."""
    if not isinstance(analysis_dict, dict):
        logging.warning("translate_analysis called with non-dictionary input.")
        return analysis_dict

    # Use language code (e.g., 'en', 'es', 'fr')
    target_language_code = target_language.lower().split('-')[0] # Get base language code
    if target_language_code == 'en':
        logging.info("Target language is English, skipping translation.")
        return analysis_dict

    logging.info(f"Attempting to translate analysis to: {target_language_code}")
    translated_analysis = {}

    try:
        # Nested function to handle translation
        def translate_value(value, dest_lang):
            if isinstance(value, str) and value.strip():
                try:
                    # Add slight delay to potentially avoid rate limits (adjust as needed)
                    # time.sleep(0.05)
                    translated = translator.translate(value, dest=dest_lang).text
                    return translated
                except Exception as trans_err:
                    logging.warning(f"Translation failed for string snippet '{value[:30]}...': {trans_err}. Returning original.")
                    return value # Return original string on error
            elif isinstance(value, list):
                # Translate each item in the list
                return [translate_value(item, dest_lang) for item in value]
            elif isinstance(value, dict):
                 # Recursively translate values in nested dictionaries
                return {k: translate_value(v, dest_lang) for k, v in value.items()}
            # Return non-string/non-list/non-dict items as is (like null, numbers, booleans)
            return value

        # Translate all values first
        for key, value in analysis_dict.items():
            translated_analysis[key] = translate_value(value, target_language_code)

        # --- Optional: Translate Keys (Uncomment if needed) ---
        # final_translated_dict = {}
        # for key, value in translated_analysis.items():
        #     try:
        #         # Replace underscores with spaces for better key translation
        #         key_to_translate = key.replace("_", " ")
        #         translated_key_text = translator.translate(key_to_translate, dest=target_language_code).text
        #         # Convert translated key back to snake_case (simple version)
        #         translated_key = translated_key_text.replace(" ", "_").lower()
        #         # Handle potential empty keys or conflicts if translation is poor
        #         if not translated_key:
        #              translated_key = key # Fallback to original key
        #         final_translated_dict[translated_key] = value
        #     except Exception as key_trans_err:
        #         logging.warning(f"Could not translate key '{key}': {key_trans_err}. Using original key.")
        #         final_translated_dict[key] = value
        # logging.info("Translation of keys and values complete.")
        # return final_translated_dict
        # --- End Optional Key Translation ---

        logging.info("Translation of values complete (keys kept original).")
        return translated_analysis # Return dict with original keys, translated values

    except Exception as e:
        logging.error(f"❌ Unexpected error during translation process: {e}", exc_info=True)
        # Return the dictionary with whatever got translated before the error
        return translated_analysis if translated_analysis else analysis_dict


# -------------------- Flask Routes --------------------

@app.route("/upload", methods=["POST"])
def upload_file():
    """Upload a file, save it, and clear previous document context AND chat history."""
    if "file" not in request.files:
        logging.warning("Upload attempt with no file part.")
        return jsonify({"error": "No file part"}), 400

    file = request.files["file"]
    if not file or file.filename == "":
        logging.warning("Upload attempt with no selected file.")
        return jsonify({"error": "No selected file"}), 400

    filename = "".join(c for c in file.filename if c.isalnum() or c in ('.', '_', '-')).strip()
    if not filename:
        filename = f"uploaded_file_{secrets.token_hex(4)}" # Add random hex if sanitization fails

    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)

    try:
        file.save(file_path)
        logging.info(f"✅ File uploaded: {file_path}")

        # --- Clear previous context AND chat history on new upload ---
        session.pop('last_analyzed_text', None)
        session.pop('last_analyzed_filename', None)
        session.pop('last_analysis_result', None)
        session.pop('chat_history', None) # <-- Clear chat history
        logging.info("Cleared previous document context, analysis, and chat history from session due to new upload.")
        session.modified = True # Ensure session changes are saved

        return jsonify({"message": "File uploaded successfully", "file_path": file_path, "filename": filename}), 200
    except Exception as e:
        logging.error(f"❌ Failed to save uploaded file {filename}: {e}", exc_info=True)
        return jsonify({"error": f"Failed to save file: {e}"}), 500


@app.route("/analyze", methods=["POST"])
def analyze_report():
    """Analyzes uploaded file (text or image) and returns structured JSON."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Request body must be JSON'}), 400

        file_path = data.get('file_path')
        filename = data.get('filename')
        language = data.get('language', 'en') or 'en' # Default to 'en'

        if not file_path or not filename:
            logging.warning("/analyze called without file_path or filename in request.")
            return jsonify({'error': 'File path and filename are required'}), 400

        # Verify file exists
        if not os.path.exists(file_path):
             logging.error(f"/analyze error: File not found at path: {file_path}")
             potential_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
             if os.path.exists(potential_path):
                 file_path = potential_path
                 logging.info(f"Corrected file path to: {file_path}")
             else:
                 return jsonify({'error': f'File not found: {filename}. It may not have been uploaded correctly or was deleted.'}), 404

        # Determine file type and call appropriate analysis function
        ext = pathlib.Path(filename).suffix.lower()
        analysis_result = None
        extracted_text_for_session = None

        logging.info(f"Analyzing file: {filename} (type: {ext}), Target language: {language}")

        # --- Image Analysis Path ---
        if ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"]:
            logging.info(f"Detected image file type ({ext}). Reading image data...")
            try:
                with open(file_path, "rb") as f:
                    image_data = f.read()
                if not image_data:
                     logging.error(f"Read 0 bytes from image file: {file_path}")
                     return jsonify({'error': f'Could not read image file content: {filename}'}), 500

                analysis_result = analyze_image_with_gemini(image_data, filename)
                # Clear any potential stale text context if analyzing an image
                session.pop('last_analyzed_text', None)
                session.pop('last_analyzed_filename', None)
                session.modified = True

            except FileNotFoundError:
                 logging.error(f"Image file not found at path during read: {file_path}")
                 return jsonify({'error': f'File not found during analysis: {filename}'}), 404
            except Exception as img_read_err:
                 logging.error(f"Error reading image file {filename}: {img_read_err}", exc_info=True)
                 return jsonify({'error': f'Error reading image file: {img_read_err}'}), 500

        # --- Text/Document Analysis Path ---
        elif ext in [".pdf", ".docx", ".xlsx", ".txt", ".rtf", ".csv"] or not ext: # Treat common docs or no extension as text-based
             logging.info(f"Detected text-based file type ({ext}). Extracting text...")
             extracted_text = extract_text(file_path)

             if not extracted_text:
                 # Provide specific feedback if OCR failed on an image file processed via extract_text
                 if os.path.splitext(filename)[1].lower() in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"] and os.path.exists(file_path):
                     logging.warning(f"OCR failed to extract text from image file via text path: {filename}")
                     return jsonify({'error': f'Could not extract text using OCR from image file: {filename}. Analysis cannot proceed.'}), 400
                 else:
                     logging.warning(f"No text could be extracted from file: {filename}")
                     return jsonify({'error': f'No text could be extracted from the file: {filename}. It might be empty, corrupted, or an unsupported format.'}), 400

             # Store extracted text for potential chat context
             extracted_text_for_session = extracted_text
             analysis_result = analyze_text_with_gemini(extracted_text)

        # --- Unsupported File Type ---
        else:
            logging.warning(f"Unsupported file type for analysis: {ext} ({filename})")
            return jsonify({'error': f'Unsupported file type: {ext}. Cannot analyze this file.'}), 400

        # --- Check Analysis Result ---
        if not analysis_result or 'error' in analysis_result:
            logging.error(f"Analysis failed for {filename}. Result: {analysis_result}")
            error_message = analysis_result.get('error', 'Unknown analysis error') if isinstance(analysis_result, dict) else "Unknown analysis error"
            # Clear potentially incomplete session data if analysis failed
            session.pop('last_analyzed_text', None)
            session.pop('last_analyzed_filename', None)
            session.pop('last_analysis_result', None)
            session.modified = True
            return jsonify({'error': error_message, 'details': analysis_result}), 500

        # --- Store context in session ONLY if text was successfully extracted and analyzed ---
        if extracted_text_for_session and analysis_result and 'error' not in analysis_result:
             session['last_analyzed_text'] = extracted_text_for_session
             session['last_analyzed_filename'] = filename
             session['last_analysis_result'] = analysis_result # Store the analysis JSON too
             logging.info(f"Stored extracted text ({len(extracted_text_for_session)} chars) and analysis result for {filename} in session.")
             session.modified = True
        # Explicitly clear if it wasn't a text analysis success
        elif 'last_analyzed_text' in session or 'last_analyzed_filename' in session:
             session.pop('last_analyzed_text', None)
             session.pop('last_analyzed_filename', None)
             session.pop('last_analysis_result', None) # Clear analysis result too if no text
             session.modified = True


        # --- Translate the successful analysis ---
        final_analysis = translate_analysis(analysis_result, language)


        # Return the final (potentially translated) analysis
        return jsonify(final_analysis), 200

    except Exception as e:
        logging.exception(f"❌ Unexpected error in /analyze route for file {data.get('filename', 'unknown')}:")
        return jsonify({'error': f'An unexpected server error occurred during analysis: {str(e)}'}), 500


@app.route("/api/chat", methods=["POST"])
def chat_api():
    """Handles chat messages, using document context from session if available."""
    user_message = request.json.get("message")
    if not user_message:
        logging.warning("Chat request with no message.")
        return jsonify({"error": "No message provided"}), 400

    # Retrieve document context from session
    document_text = session.get('last_analyzed_text')
    filename = session.get('last_analyzed_filename')

    prompt_for_gemini = ""
    max_context_chars = 50000 # Limit context sent to Gemini to avoid hitting token limits aggressively

    try:
        if document_text and filename:
            logging.info(f"Chat query received. Using context from document: {filename}")
            # Truncate context if it's very large
            truncated_context = document_text
            if len(document_text) > max_context_chars:
                logging.warning(f"Document context length ({len(document_text)}) exceeds limit ({max_context_chars}), truncating for chat prompt.")
                truncated_context = document_text[:max_context_chars] + "\n... [Content Truncated]"


            # Construct prompt telling Gemini to use the context if relevant
            prompt_for_gemini = f"""Forget that you are an ai and imagine your self as a helpful medical AI assistant. You have access to the content of a medical document named '{filename}'. Also dont I am an AI assistant and cannot provide medical advice instead tell then that i have suggested the medicine but only take it under a doctor supervision.

            *Instructions:*
            1. Answer the user's query below.
            2. *Prioritize using the provided document context* if the query seems related to it. Clearly state if your answer comes from the document.
            3. If the query is not related to the document context, answer it using your general medical knowledge.
            4. If you are unsure or the document doesn't contain the answer, state that clearly.
            5. Respond concisely and clearly. Use bullet points if appropriate.
            6.Dont tell that I am an AI and cannot provide medical advice. 

            *Document Context ('{filename}'):*
            --- START CONTEXT ---
            {truncated_context}
            --- END CONTEXT ---

            *User Query:* {user_message}
            """
        else:
            logging.info("Chat query received. No document context in session.")
            # Construct a general prompt
            prompt_for_gemini = f"""You are a helpful medical AI assistant.
            Answer the following user query using your general medical knowledge.
            Respond concisely and clearly. Use bullet points if appropriate.

            *User Query:* {user_message}
            """

        # Send the constructed prompt to Gemini
        # Use generate_content for a single-turn response based on the prompt
        logging.debug(f"Sending prompt to Gemini for chat (length: {len(prompt_for_gemini)} chars)")
        response = gemini_model.generate_content(prompt_for_gemini)

        # Process the response
        if hasattr(response, "text"):
            response_text = response.text
            logging.info("Received chat response from Gemini.")
            return jsonify({"response": response_text})
        elif hasattr(response, 'candidates') and response.candidates:
             candidate = response.candidates[0]
             if candidate.finish_reason != 'STOP':
                  logging.error(f"❌ Gemini chat generation stopped unexpectedly. Reason: {candidate.finish_reason}")
                  safety_ratings = getattr(candidate, 'safety_ratings', 'N/A')
                  return jsonify({"error": f"Chat generation failed or was blocked. Reason: {candidate.finish_reason}", "safety_ratings": safety_ratings}), 500
             else:
                  # Try to extract text if available
                  try:
                      response_text = candidate.content.parts[0].text.strip()
                      logging.warning("Gemini chat response had no top-level .text, but found text in candidate.")
                      return jsonify({"response": response_text})
                  except (AttributeError, IndexError):
                       logging.error("❌ No 'text' attribute and couldn't extract chat response from candidate parts.")
                       return jsonify({"error": "No text content received from Gemini for chat."}), 500
        else:
            logging.error("❌ No 'text' or valid 'candidates' in Gemini chat response.")
            logging.debug(f"Full Gemini Chat Response: {response}")
            return jsonify({"error": "No text content received from Gemini for chat."}), 500

    except Exception as e:
        logging.exception("❌ Error during chat processing:") # Log full traceback
        return jsonify({"error": f"An unexpected error occurred during chat: {e}"}), 500

@app.route('/api/auth/signup', methods=['POST'])
def signup():
    data = request.get_json()
    
    if not data or not data.get('email') or not data.get('password'):
        return jsonify({'message': 'Missing email or password'}), 400
    
    try:
        # Check if user already exists
        users_ref = db.collection('users')
        existing_user = users_ref.where('email', '==', data['email']).get()
        
        if len(list(existing_user)) > 0:
            return jsonify({'message': 'User already exists'}), 400
        
        # Create new user
        new_user = {
            'email': data['email'],
            'password': generate_password_hash(data['password']),
            'name': data.get('name', ''),
            'avatar': None,
            'role': 'user',
            'personalInfo': {
                'dateOfBirth': None,
                'age': None,
                'gender': None,
                'bloodType': None
            },
            'contactInfo': {
                'phone': None,
                'address': None,
                'emergencyContact': None
            },
            'created_at': firestore.SERVER_TIMESTAMP
        }
        
        # Add user to Firestore
        user_ref = users_ref.document()
        user_ref.set(new_user)
        
        # Generate token
        token = jwt.encode({
            'user_id': user_ref.id,
            'exp': datetime.datetime.utcnow() + datetime.timedelta(days=1)
        }, JWT_SECRET_KEY)
        
        return jsonify({
            'message': 'User created successfully',
            'token': token,
            'user_id': user_ref.id
        }), 201
        
    except Exception as e:
        logging.error(f"Error in signup: {str(e)}")
        return jsonify({'message': 'Error creating user'}), 500

@app.route('/api/auth/signin', methods=['POST'])
def signin():
    data = request.get_json()
    
    if not data or not data.get('email') or not data.get('password'):
        return jsonify({'message': 'Missing email or password'}), 400
    
    try:
        # Find user by email
        users_ref = db.collection('users')
        user_docs = users_ref.where('email', '==', data['email']).get()
        
        user_list = list(user_docs)
        if not user_list:
            return jsonify({'message': 'User not found'}), 401
        
        user_doc = user_list[0]
        user_data = user_doc.to_dict()
        
        # Check password
        if not check_password_hash(user_data['password'], data['password']):
            return jsonify({'message': 'Invalid password'}), 401
        
        # Generate token
        token = jwt.encode({
            'user_id': user_doc.id,
            'exp': datetime.datetime.utcnow() + datetime.timedelta(days=1)
        }, JWT_SECRET_KEY)
        
        return jsonify({
            'message': 'Logged in successfully',
            'token': token,
            'user_id': user_doc.id
        }), 200
        
    except Exception as e:
        logging.error(f"Error in signin: {str(e)}")
        return jsonify({'message': 'Error during signin'}), 500

# -------------------- Run the Flask App --------------------

if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    debug_mode = os.getenv("FLASK_DEBUG", "True").lower() == "true"
    logging.info(f"Starting Flask server on port {port} with debug mode: {debug_mode}")
    # Use host='0.0.0.0' to make it accessible on the local network
    # For production, use a WSGI server like Gunicorn or Waitress and set debug=False
    app.run(debug=debug_mode, host='0.0.0.0', port=port)